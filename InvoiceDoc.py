'''
Created on 2014-4-22

@author: Sindo
'''

import os,sys,locale
import re

from docx import *
from datetime import date
from datetime import datetime
from datetime import timedelta

from Navigator import Navigator

def printSecSepor(seporName="========") :
	print ""
	print "========================================================================"
	print "=========================== %s ===================================" % (seporName)
	print "========================================================================"
	print ""
	
class InvoiceDoc :
	def __init__(self, userName, userPswd, workDir, debug=False):
		self.debug = debug
		
		self.userName = userName
		self.userPswd = userPswd
		
		self.workDir = workDir

	def perform(self, docFile):
		print "docFile is " + (docFile)
		
		
		document = Document(docFile)
		sections = document.sections
		
		print "section_length:%d" % len(sections)

	def testCase(self):
		document = Document()

		document.add_heading(u'云水谣主题餐厅', 0)

		p = document.add_paragraph(u'小河直街店')
		p.add_run('bold').bold = True
		p.add_run(' and some ')
		p.add_run('italic.').italic = True

		document.add_heading('Heading, level 1', level=1)
		document.add_paragraph('Intense quote', style='IntenseQuote')

		document.add_paragraph(
			'first item in unordered list', style='ListBullet'
		)
		document.add_paragraph(
			'first item in ordered list', style='ListNumber'
		)

		#document.add_picture('monty-truth.png', width=Inches(1.25))

		recordset = []
		
		infoColDict = {
				"qty" : 0,
				"id": 1, 
				"desc": "TEST"
		}
		recordset.append(infoColDict)
		
		infoColDict = {
				"qty" : 0,
				"id": 1, 
				"desc": "TEST"
		}
		recordset.append(infoColDict)
		
		infoColDict = {
				"qty" : 0,
				"id": 1, 
				"desc": "TEST"
		}
		recordset.append(infoColDict)
		

		table = document.add_table(rows=1, cols=3)
		hdr_cells = table.rows[0].cells
		hdr_cells[0].text = 'Qty'
		hdr_cells[1].text = 'Id'
		hdr_cells[2].text = 'Desc'
		for item in recordset:
			row_cells = table.add_row().cells
			row_cells[0].text = str(item["qty"])
			row_cells[1].text = str(item["id"])
			row_cells[2].text = item["desc"]

		document.add_page_break()

		document.save('demo.docx')

if __name__ == '__main__':	# sys.argv[0], sys.argv[1]
	reload(sys)
	sys.setdefaultencoding("utf-8")
	
	print "InvoiceDoc_process_begin......"  + sys.argv[0]
	
	resVal = False
	debugMode = False
	
	userName = 'lijj'
	userPswd = os.getenv("aintPswd")
	
	workDir = u"E:"
	
	invoiceDoc = InvoiceDoc(userName, userPswd, workDir, debugMode)
	
	docFile = u"InvoiceTest.docx"
	
	resVal = invoiceDoc.testCase()
	resVal= invoiceDoc.perform(docFile)
	
	printSecSepor()
	
	if resVal: print "congratulations_your_work_is_done_successfully" 
	else:      print "unfortunately_your_work_is_failed"